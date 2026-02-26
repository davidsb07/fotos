from __future__ import annotations

import tempfile
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageOps


EXIF_ORIENTATION_TAG = 274
ORIENTACOES_ROTACIONADAS = {5, 6, 7, 8}
ORIENTACOES_TRANSPOSTAS = {2, 3, 4, 5, 6, 7, 8}


@dataclass(slots=True)
class FotoDocumento:
    caminho: Path
    descricao: str


def _obter_dimensoes_imagem(caminho: Path) -> tuple[float, float]:
    try:
        with Image.open(caminho) as img:
            largura, altura = img.size
            exif = img.getexif()
            orientacao = exif.get(EXIF_ORIENTATION_TAG) if exif else None
            if orientacao in ORIENTACOES_ROTACIONADAS:
                largura, altura = altura, largura
            return float(largura), float(altura)
    except Exception:
        return (4.0, 3.0)


def _corrigir_orientacao_exif(caminho: Path, temporarios: list[Path]) -> Path:
    try:
        with Image.open(caminho) as img:
            exif = img.getexif()
            orientacao = exif.get(EXIF_ORIENTATION_TAG) if exif else None
            if orientacao not in ORIENTACOES_TRANSPOSTAS:
                return caminho

            imagem_corrigida = ImageOps.exif_transpose(img)
            arquivo_tmp = Path(tempfile.gettempdir()) / f"corrigido_{uuid.uuid4().hex}_{caminho.name}"

            save_kwargs: dict[str, int] = {}
            if (img.format or "").upper() in {"JPEG", "JPG"}:
                save_kwargs["quality"] = 95

            imagem_corrigida.save(arquivo_tmp, **save_kwargs)
            temporarios.append(arquivo_tmp)
            return arquivo_tmp
    except Exception:
        return caminho


def _eh_vertical(caminho: Path) -> bool:
    largura, altura = _obter_dimensoes_imagem(caminho)
    return altura > largura


def _calcular_dimensoes_foto(caminho: Path, largura_max: float, altura_max: float) -> tuple[float, float]:
    largura_px, altura_px = _obter_dimensoes_imagem(caminho)
    proporcao = largura_px / altura_px

    if proporcao >= 1:
        largura = largura_max
        altura = largura / proporcao
        if altura > altura_max:
            altura = altura_max
            largura = altura * proporcao
    else:
        altura = altura_max
        largura = altura * proporcao
        if largura > largura_max:
            largura = largura_max
            altura = largura / proporcao

    return largura, altura


def criar_documento_fotos(
    titulo: str,
    origem_fotos: str,
    fotos: list[FotoDocumento],
    arquivo_saida: Path,
    largura_coluna: float = 8.5,
    altura_padrao: float = 6.4,
) -> None:
    if not fotos:
        raise ValueError("A lista de fotos nao pode ser vazia.")

    arquivo_saida.parent.mkdir(parents=True, exist_ok=True)
    documento = Document()
    temporarios: list[Path] = []

    try:
        for secao in documento.sections:
            secao.top_margin = Cm(2)
            secao.bottom_margin = Cm(2)
            secao.left_margin = Cm(2)
            secao.right_margin = Cm(2)

        p_titulo = documento.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run(titulo)
        run_titulo.bold = True
        run_titulo.font.size = Pt(10)
        run_titulo.font.name = "Arial"

        p_origem = documento.add_paragraph()
        p_origem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_origem = p_origem.add_run(origem_fotos)
        run_origem.bold = True
        run_origem.font.size = Pt(9)
        run_origem.font.name = "Arial"

        documento.add_paragraph()

        num_fotos = len(fotos)
        num_pares = (num_fotos + 1) // 2
        tabela = documento.add_table(rows=num_pares * 2, cols=2)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(0, num_fotos, 2):
            linha_foto = (i // 2) * 2
            linha_desc = linha_foto + 1

            foto_1 = fotos[i]
            foto_2 = fotos[i + 1] if i + 1 < num_fotos else None

            caminho_1 = _corrigir_orientacao_exif(foto_1.caminho, temporarios)
            caminho_2 = _corrigir_orientacao_exif(foto_2.caminho, temporarios) if foto_2 else None

            foto_1_vertical = _eh_vertical(caminho_1)
            foto_2_vertical = bool(caminho_2 and _eh_vertical(caminho_2))
            ambas_verticais = foto_1_vertical and foto_2_vertical
            verticais_especial = ambas_verticais or (foto_1_vertical and foto_2 is None)

            altura_max = 12.0 if verticais_especial else altura_padrao
            largura_col = largura_coluna * 0.70 if verticais_especial else largura_coluna

            celula_foto_1 = tabela.cell(linha_foto, 0)
            celula_desc_1 = tabela.cell(linha_desc, 0)

            larg_1, alt_1 = _calcular_dimensoes_foto(caminho_1, largura_col, altura_max)
            p_img_1 = celula_foto_1.paragraphs[0]
            p_img_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_1 = p_img_1.add_run()
            run_img_1.add_picture(str(caminho_1), width=Cm(larg_1), height=Cm(alt_1))

            p_desc_1 = celula_desc_1.paragraphs[0]
            p_desc_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_desc_1 = p_desc_1.add_run(foto_1.descricao)
            run_desc_1.font.size = Pt(9)
            run_desc_1.font.name = "Arial"

            if foto_2 and caminho_2:
                celula_foto_2 = tabela.cell(linha_foto, 1)
                celula_desc_2 = tabela.cell(linha_desc, 1)

                larg_2, alt_2 = _calcular_dimensoes_foto(caminho_2, largura_col, altura_max)
                p_img_2 = celula_foto_2.paragraphs[0]
                p_img_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img_2 = p_img_2.add_run()
                run_img_2.add_picture(str(caminho_2), width=Cm(larg_2), height=Cm(alt_2))

                p_desc_2 = celula_desc_2.paragraphs[0]
                p_desc_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_desc_2 = p_desc_2.add_run(foto_2.descricao)
                run_desc_2.font.size = Pt(9)
                run_desc_2.font.name = "Arial"

        documento.save(str(arquivo_saida))
    finally:
        for arquivo_tmp in temporarios:
            try:
                arquivo_tmp.unlink(missing_ok=True)
            except Exception:
                continue
